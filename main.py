
from fastapi import FastAPI, Request
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from docx import Document
import os
from datetime import datetime

app = FastAPI()

# Certifique-se que o diretório de laudos existe
os.makedirs("static/laudos", exist_ok=True)

app.mount("/static", StaticFiles(directory="static"), name="static")

class LaudoRequest(BaseModel):
    paciente: str
    data: str
    corpo: str
    conclusao: str
    solicitante: str = None

@app.post("/gerar-laudo")
def gerar_laudo(req: LaudoRequest):
    try:
        document = Document()
        document.add_heading('Laudo Ultrassonográfico', 0)

        document.add_paragraph(f"Paciente: {req.paciente}")
        if req.solicitante:
            document.add_paragraph(f"Médico Solicitante: {req.solicitante}")
        document.add_paragraph(f"Data do Exame: {req.data}")

        document.add_heading("Achados", level=1)
        document.add_paragraph(req.corpo)

        document.add_heading("Conclusão", level=1)
        document.add_paragraph(req.conclusao)

        filename = f"Laudo_{req.paciente.replace(' ', '_')}.docx"
        output_path = f"static/laudos/{filename}"
        document.save(output_path)

        return {"url": f"https://apidocx.onrender.com/static/laudos/{filename}"}
    except Exception as e:
        return JSONResponse(status_code=500, content={"detail": str(e)})
